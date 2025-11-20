#!/usr/bin/env python3
"""
GenieFactory BP 14 Mois - Script 1: Extraction de Données
Parse BP Excel, BM Word, Pacte → JSON structuré

Inputs:
  - data/raw/BP_FABRIQ_PRODUCT-OCT2025.xlsx
  - data/raw/Business_Plan_GenieFactory-SEPT2025.docx
  - data/raw/GENIE_FACTORY_PACTE_AATL-v3 [1].docx

Outputs:
  - data/structured/bp_extracted.json
  - data/structured/bm_extracted.json
  - data/structured/pacte_extracted.json
"""

import json
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List

import openpyxl
from openpyxl.cell import Cell
from docx import Document

# Configuration logging
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)


class BPExcelExtractor:
    """Extracteur pour le Business Plan Excel"""

    def __init__(self, filepath: Path):
        self.filepath = filepath
        self.workbook = None

    def load(self):
        """Charger le workbook Excel"""
        logger.info(f"📂 Chargement BP Excel: {self.filepath}")
        self.workbook = openpyxl.load_workbook(self.filepath, data_only=False)
        logger.info(f"✓ Workbook chargé - {len(self.workbook.sheetnames)} sheets")

    def get_cell_value(self, cell: Cell) -> Any:
        """Extraire valeur ou formule d'une cellule"""
        if cell.value is None:
            return None

        # Si c'est une formule, on garde les deux
        if hasattr(cell, 'data_type') and cell.data_type == 'f':
            return {
                'type': 'formula',
                'formula': str(cell.value),
                'value': None  # Sera calculé par Excel
            }

        return cell.value

    def extract_pricing(self) -> Dict[str, Any]:
        """Extraire pricing depuis sheet Paramètres"""
        logger.info("📊 Extraction pricing...")

        pricing = {}

        # Essayer de trouver la sheet Paramètres
        sheet_names = [s.lower() for s in self.workbook.sheetnames]
        param_sheet = None

        for sheet_name in self.workbook.sheetnames:
            if 'param' in sheet_name.lower() or 'price' in sheet_name.lower():
                param_sheet = self.workbook[sheet_name]
                break

        if not param_sheet:
            logger.warning("⚠️ Sheet Paramètres non trouvée, utilisation valeurs par défaut")
            return {
                'hackathon': {'m1_m6': 18000, 'm7_m14': 20000},
                'factory': {'m1_m6': 75000, 'm7_m14': 82500},
                'services': {'m1_m6': 10000, 'm7_m14': 17500},
                'formation': {'m1_m6': 5000, 'm7_m14': 5500}
            }

        # Scanner les premières lignes pour trouver les prix
        logger.info(f"✓ Sheet trouvée: {param_sheet.title}")

        # Extraction basique - chercher patterns de prix
        for row_idx in range(1, min(30, param_sheet.max_row + 1)):
            for col_idx in range(1, min(10, param_sheet.max_column + 1)):
                cell = param_sheet.cell(row_idx, col_idx)
                value = self.get_cell_value(cell)

                if value and isinstance(value, (int, float)):
                    # Détecter pricing typique (entre 5K et 100K)
                    if 5000 <= value <= 100000:
                        # Chercher label dans colonne précédente
                        label_cell = param_sheet.cell(row_idx, max(1, col_idx - 1))
                        label = str(label_cell.value or '').lower()

                        if 'hackathon' in label or 'hack' in label:
                            pricing.setdefault('hackathon', {})['price'] = value
                        elif 'factory' in label:
                            pricing.setdefault('factory', {})['price'] = value
                        elif 'service' in label or 'impl' in label:
                            pricing.setdefault('services', {})['price'] = value

        logger.info(f"✓ Pricing extrait: {len(pricing)} offres")
        return pricing

    def extract_pl_data(self) -> Dict[str, Any]:
        """Extraire données P&L"""
        logger.info("📊 Extraction P&L...")

        pl_data = {
            'revenue_lines': [],
            'cost_lines': [],
            'months': []
        }

        # Chercher sheet P&L
        pl_sheet = None
        for sheet_name in self.workbook.sheetnames:
            if 'p&l' in sheet_name.lower() or 'p&amp;l' in sheet_name.lower() or 'compte' in sheet_name.lower():
                pl_sheet = self.workbook[sheet_name]
                break

        if not pl_sheet:
            logger.warning("⚠️ Sheet P&L non trouvée")
            return pl_data

        logger.info(f"✓ Sheet P&L trouvée: {pl_sheet.title}")

        # Scanner les données
        for row_idx in range(1, min(50, pl_sheet.max_row + 1)):
            row_label = pl_sheet.cell(row_idx, 1).value
            if row_label:
                row_label_str = str(row_label).lower()

                # Identifier les lignes de revenus
                if any(keyword in row_label_str for keyword in ['hackathon', 'factory', 'hub', 'service', 'ca', 'chiffre']):
                    row_data = {
                        'label': str(row_label),
                        'values': []
                    }

                    # Extraire valeurs mensuelles (colonnes 2+)
                    for col_idx in range(2, min(20, pl_sheet.max_column + 1)):
                        cell_value = self.get_cell_value(pl_sheet.cell(row_idx, col_idx))
                        if isinstance(cell_value, (int, float)):
                            row_data['values'].append(cell_value)

                    if row_data['values']:
                        pl_data['revenue_lines'].append(row_data)

        logger.info(f"✓ P&L extrait: {len(pl_data['revenue_lines'])} lignes revenus")
        return pl_data

    def extract_all(self) -> Dict[str, Any]:
        """Extraction complète"""
        self.load()

        data = {
            'meta': {
                'source_file': self.filepath.name,
                'extracted_at': datetime.now().isoformat(),
                'sheets': self.workbook.sheetnames
            },
            'pricing': self.extract_pricing(),
            'pl_data': self.extract_pl_data()
        }

        return data


class WordExtractor:
    """Extracteur pour documents Word"""

    def __init__(self, filepath: Path):
        self.filepath = filepath
        self.document = None

    def load(self):
        """Charger le document Word"""
        logger.info(f"📂 Chargement Word: {self.filepath}")
        self.document = Document(self.filepath)
        logger.info(f"✓ Document chargé - {len(self.document.paragraphs)} paragraphes")

    def extract_tables(self) -> List[Dict[str, Any]]:
        """Extraire tous les tableaux"""
        logger.info("📊 Extraction tableaux...")

        tables_data = []
        for idx, table in enumerate(self.document.tables):
            table_data = {
                'index': idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'content': []
            }

            # Extraire contenu
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data['content'].append(row_data)

            tables_data.append(table_data)

        logger.info(f"✓ {len(tables_data)} tableaux extraits")
        return tables_data

    def extract_financial_metrics(self) -> Dict[str, Any]:
        """Extraire métriques financières depuis texte"""
        logger.info("📊 Extraction métriques financières...")

        metrics = {}
        full_text = '\n'.join([p.text for p in self.document.paragraphs])

        # Pattern ARR
        arr_matches = re.findall(r'ARR[:\s]*(\d+)[\s]?K€', full_text, re.IGNORECASE)
        if arr_matches:
            metrics['arr_values'] = [int(m) * 1000 for m in arr_matches]
            logger.info(f"✓ ARR trouvé: {metrics['arr_values']}")

        # Pattern CA
        ca_matches = re.findall(r'CA[:\s]*(\d+)[\s]?K€', full_text, re.IGNORECASE)
        if ca_matches:
            metrics['ca_values'] = [int(m) * 1000 for m in ca_matches]

        # Pattern conversion
        conv_matches = re.findall(r'conversion[:\s]*(\d+)%', full_text, re.IGNORECASE)
        if conv_matches:
            metrics['conversion_rates'] = [int(m) / 100 for m in conv_matches]

        return metrics

    def extract_all(self) -> Dict[str, Any]:
        """Extraction complète"""
        self.load()

        data = {
            'meta': {
                'source_file': self.filepath.name,
                'extracted_at': datetime.now().isoformat(),
                'paragraphs_count': len(self.document.paragraphs),
                'tables_count': len(self.document.tables)
            },
            'tables': self.extract_tables(),
            'metrics': self.extract_financial_metrics()
        }

        return data


class PacteExtractor(WordExtractor):
    """Extracteur spécialisé pour le Pacte Actionnaires"""

    def extract_arr_milestones(self) -> List[Dict[str, Any]]:
        """Extraire milestones ARR depuis pacte"""
        logger.info("📊 Extraction milestones ARR...")

        milestones = []
        full_text = '\n'.join([p.text for p in self.document.paragraphs])

        # Patterns milestones ARR
        patterns = [
            r'ARR\s*[≥>=]\s*(\d+)[\s]?K€',
            r'ARR\s*[≥>=]\s*(\d[\s,.]?\d+)[\s]?M€',
            r'(\d+)[\s]?K€.*ARR',
            r'(\d[\s,.]?\d+)[\s]?M€.*ARR'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, full_text, re.IGNORECASE)
            for match in matches:
                # Nettoyer le match
                value_str = match.replace(' ', '').replace(',', '').replace('.', '')

                try:
                    value = int(value_str)
                    # Si pattern en M€
                    if 'M€' in pattern:
                        value *= 1000000
                    elif 'K€' in pattern:
                        value *= 1000

                    if value > 100000:  # Filtre valeurs réalistes ARR
                        milestones.append({
                            'arr_target': value,
                            'pattern': pattern
                        })
                except ValueError:
                    continue

        # Déduplication
        unique_milestones = []
        seen = set()
        for m in milestones:
            if m['arr_target'] not in seen:
                seen.add(m['arr_target'])
                unique_milestones.append(m)

        logger.info(f"✓ {len(unique_milestones)} milestones ARR trouvés: {[m['arr_target'] for m in unique_milestones]}")
        return unique_milestones

    def extract_all(self) -> Dict[str, Any]:
        """Extraction complète pacte"""
        self.load()

        data = {
            'meta': {
                'source_file': self.filepath.name,
                'extracted_at': datetime.now().isoformat()
            },
            'arr_milestones': self.extract_arr_milestones(),
            'metrics': self.extract_financial_metrics()
        }

        return data


def main():
    """Fonction principale"""
    logger.info("="*60)
    logger.info("🚀 EXTRACTION DONNÉES - GenieFactory BP 14 Mois")
    logger.info("="*60)

    # Chemins fichiers
    base_path = Path(__file__).parent.parent
    bp_excel_path = base_path / "data" / "raw" / "BP FABRIQ_PRODUCT-OCT2025.xlsx"
    bm_word_path = base_path / "data" / "raw" / "Business Plan GenieFactory-SEPT2025.docx"
    pacte_word_path = base_path / "data" / "raw" / "GENIE FACTORY PACTE AATL-v3 [1].docx"

    output_dir = base_path / "data" / "structured"
    output_dir.mkdir(parents=True, exist_ok=True)

    # 1. Extraction BP Excel
    logger.info("\n📊 Phase 1: Extraction BP Excel")
    bp_extractor = BPExcelExtractor(bp_excel_path)
    bp_data = bp_extractor.extract_all()

    bp_output = output_dir / "bp_extracted.json"
    with open(bp_output, 'w', encoding='utf-8') as f:
        json.dump(bp_data, f, indent=2, ensure_ascii=False)
    logger.info(f"✅ BP Excel extrait → {bp_output}")

    # 2. Extraction BM Word
    logger.info("\n📄 Phase 2: Extraction BM Word")
    bm_extractor = WordExtractor(bm_word_path)
    bm_data = bm_extractor.extract_all()

    bm_output = output_dir / "bm_extracted.json"
    with open(bm_output, 'w', encoding='utf-8') as f:
        json.dump(bm_data, f, indent=2, ensure_ascii=False)
    logger.info(f"✅ BM Word extrait → {bm_output}")

    # 3. Extraction Pacte
    logger.info("\n📜 Phase 3: Extraction Pacte Actionnaires")
    pacte_extractor = PacteExtractor(pacte_word_path)
    pacte_data = pacte_extractor.extract_all()

    pacte_output = output_dir / "pacte_extracted.json"
    with open(pacte_output, 'w', encoding='utf-8') as f:
        json.dump(pacte_data, f, indent=2, ensure_ascii=False)
    logger.info(f"✅ Pacte extrait → {pacte_output}")

    # Résumé
    logger.info("\n" + "="*60)
    logger.info("✅ EXTRACTION TERMINÉE")
    logger.info("="*60)
    logger.info(f"📁 Fichiers générés:")
    logger.info(f"  • {bp_output}")
    logger.info(f"  • {bm_output}")
    logger.info(f"  • {pacte_output}")

    # Validation basique
    logger.info("\n🔍 Validation basique:")

    pricing_ok = len(bp_data.get('pricing', {})) > 0
    logger.info(f"  {'✓' if pricing_ok else '✗'} Pricing extrait: {len(bp_data.get('pricing', {}))} offres")

    tables_ok = len(bm_data.get('tables', [])) >= 2
    logger.info(f"  {'✓' if tables_ok else '✗'} BM tableaux: {len(bm_data.get('tables', []))} tableaux")

    milestones_ok = len(pacte_data.get('arr_milestones', [])) > 0
    logger.info(f"  {'✓' if milestones_ok else '✗'} ARR milestones: {len(pacte_data.get('arr_milestones', []))}")

    if pricing_ok and tables_ok and milestones_ok:
        logger.info("\n✅ Validation: OK - Prêt pour Phase 2 (génération assumptions)")
    else:
        logger.warning("\n⚠️ Validation: Certaines données manquantes - vérifier manuellement")

    return 0


if __name__ == "__main__":
    exit(main())
