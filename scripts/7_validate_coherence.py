#!/usr/bin/env python3
"""
GenieFactory BP 14 Mois - Script 7: Validation Cohérence Avancée
Détecte les incohérences critiques dans le Business Plan

Input:
  - data/structured/projections.json
  - data/structured/assumptions.yaml
  - data/validation_rules.yaml
  - data/outputs/BM_Updated_14M.docx

Output:
  - logs/coherence_report_{timestamp}.txt
  - data/structured/corrections_proposed.yaml
"""

import json
import yaml
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Tuple
from docx import Document
from rich.console import Console
from rich.table import Table
from rich.panel import Panel

# Configuration logging
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)
console = Console()


class CoherenceValidator:
    """Validateur de cohérence avancé pour Business Plan"""

    def __init__(self, projections: List[Dict], assumptions: Dict, rules: Dict):
        self.projections = projections
        self.assumptions = assumptions
        self.rules = rules
        self.errors = []
        self.warnings = []
        self.corrections = []

    def validate_all(self, word_doc_path: Path) -> Dict[str, Any]:
        """Exécuter toutes les validations"""
        console.print("\n[bold cyan]🔍 VALIDATION COHÉRENCE AVANCÉE[/]")
        console.print("=" * 60)

        # 1. Extraire données du Word
        doc_data = self.extract_doc_data(word_doc_path)

        # 2. Valider valorisation vs ARR
        self.validate_valuation(doc_data)

        # 3. Valider cohérence inter-sections
        self.validate_inter_sections(doc_data)

        # 4. Détecter red flags
        self.validate_red_flags()

        # 5. Vérifier patterns d'erreurs
        self.check_error_patterns(word_doc_path)

        # 6. Générer rapport
        report = self.generate_report()

        return report

    def extract_doc_data(self, doc_path: Path) -> Dict[str, Any]:
        """Extraire données clés du document Word"""
        console.print("\n[cyan]📄 EXTRACTION DONNÉES DOCUMENT[/]")

        doc = Document(doc_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])

        data = {
            'full_text': full_text,
            'valorisation_mentions': [],
            'arr_mentions': [],
            'effectifs_mentions': [],
            'ca_mentions': []
        }

        # Extraire valorisations (ex: "200M€", "200-300M€")
        val_pattern = r'(\d+)(?:-(\d+))?M€'
        for match in re.finditer(val_pattern, full_text):
            val_min = int(match.group(1)) * 1_000_000
            val_max = int(match.group(2)) * 1_000_000 if match.group(2) else val_min
            context = full_text[max(0, match.start()-50):min(len(full_text), match.end()+50)]

            if any(keyword in context.lower() for keyword in ['valorisation', 'valuation', 'valued at']):
                data['valorisation_mentions'].append({
                    'min': val_min,
                    'max': val_max,
                    'text': match.group(0),
                    'context': context.strip()
                })

        # Extraire ARR (ex: "ARR 827K€", "ARR de 6.75M€")
        arr_pattern = r'ARR.*?(\d[\d\s,\.]+)\s*([KM])€'
        for match in re.finditer(arr_pattern, full_text, re.IGNORECASE):
            value_str = match.group(1).replace(' ', '').replace(',', '').replace('.', '')
            multiplier = 1000 if match.group(2) == 'K' else 1_000_000
            arr_value = int(value_str) * multiplier

            data['arr_mentions'].append({
                'value': arr_value,
                'text': match.group(0)
            })

        logger.info(f"✓ Extractions: {len(data['valorisation_mentions'])} valorisations, {len(data['arr_mentions'])} ARR")
        return data

    def validate_valuation(self, doc_data: Dict[str, Any]):
        """Valider cohérence valorisation vs ARR"""
        console.print("\n[cyan]💰 VALIDATION VALORISATION VS ARR[/]")

        # ARR M14 depuis projections
        arr_m14 = self.projections[13]['metrics']['arr']

        # Multiples de référence
        multiples = self.rules['valuation_multiples']

        # Vérifier chaque mention de valorisation
        for val_mention in doc_data['valorisation_mentions']:
            val_min = val_mention['min']
            val_max = val_mention['max']
            val_avg = (val_min + val_max) / 2

            # Calculer multiple implicite
            if arr_m14 > 0:
                multiple_min = val_min / arr_m14
                multiple_max = val_max / arr_m14
                multiple_avg = val_avg / arr_m14
            else:
                multiple_avg = 0

            # Déterminer si le multiple est cohérent
            is_realistic = multiples['realistic']['min'] <= multiple_avg <= multiples['realistic']['max']
            is_aggressive = multiples['aggressive']['min'] <= multiple_avg <= multiples['aggressive']['max']
            is_unrealistic = multiple_avg > multiples['unrealistic']['threshold']

            if is_unrealistic:
                # ERREUR CRITIQUE
                self.errors.append({
                    'type': 'INCOHÉRENCE CRITIQUE',
                    'section': '1.3 Vision',
                    'donnee_actuelle': f"{val_min/1_000_000:.0f}-{val_max/1_000_000:.0f}M€",
                    'probleme': f"Multiple {multiple_avg:.1f}x hors norme SaaS (standard 7-10x)",
                    'arr_reference': f"{arr_m14/1_000_000:.2f}M€",
                    'impact': 5,
                    'context': val_mention['context']
                })

                # Proposer corrections
                val_conservative = arr_m14 * multiples['conservative']['max']
                val_realistic = arr_m14 * multiples['realistic']['max']
                val_aggressive = arr_m14 * multiples['aggressive']['max']

                self.corrections.append({
                    'section': '1.3 Vision',
                    'field': 'valorisation_2028',
                    'options': {
                        'conservative': {
                            'value': f"{val_conservative/1_000_000:.0f}M€",
                            'multiple': multiples['conservative']['max'],
                            'justification': multiples['conservative']['context']
                        },
                        'realistic': {
                            'value': f"{val_realistic/1_000_000:.0f}M€",
                            'multiple': multiples['realistic']['max'],
                            'justification': multiples['realistic']['context']
                        },
                        'aggressive': {
                            'value': f"{val_aggressive/1_000_000:.0f}M€",
                            'multiple': multiples['aggressive']['max'],
                            'justification': multiples['aggressive']['context']
                        }
                    },
                    'recommendation': 'realistic'
                })

                console.print(f"  [red]✗ Valorisation {val_mention['text']}: Multiple {multiple_avg:.1f}x INCOHÉRENT[/]")
                console.print(f"    ARR M14: {arr_m14/1_000:,.0f}K€")
                console.print(f"    Valorisation réaliste: {val_realistic/1_000_000:.0f}M€ (10x)")

            elif is_aggressive:
                self.warnings.append({
                    'type': 'Valorisation agressive',
                    'message': f"Multiple {multiple_avg:.1f}x nécessite justification croissance >100%/an",
                    'section': '1.3 Vision'
                })
                console.print(f"  [yellow]⚠ Valorisation {val_mention['text']}: Multiple {multiple_avg:.1f}x AGRESSIF[/]")

            else:
                console.print(f"  [green]✓ Valorisation {val_mention['text']}: Multiple {multiple_avg:.1f}x cohérent[/]")

    def validate_inter_sections(self, doc_data: Dict[str, Any]):
        """Valider cohérence entre sections"""
        console.print("\n[cyan]🔗 VALIDATION COHÉRENCE INTER-SECTIONS[/]")

        # Check 1: CA total vs détail revenus
        ca_total_proj = sum(p['revenue']['total'] for p in self.projections)

        ca_detail = sum([
            sum(p['revenue']['hackathon']['revenue'] for p in self.projections),
            sum(p['revenue']['factory']['revenue'] for p in self.projections),
            sum(p['revenue']['enterprise_hub']['mrr'] for p in self.projections),
            sum(p['revenue']['services']['revenue'] for p in self.projections)
        ])

        if abs(ca_total_proj - ca_detail) / ca_total_proj > 0.01:
            self.errors.append({
                'type': 'INCOHÉRENCE INTERNE',
                'section': '7.2 Projections',
                'probleme': f"CA total ({ca_total_proj/1000:.0f}K€) != somme détails ({ca_detail/1000:.0f}K€)",
                'impact': 4
            })
            console.print(f"  [red]✗ Incohérence CA total vs détail[/]")
        else:
            console.print(f"  [green]✓ CA total cohérent avec détail revenus[/]")

    def validate_red_flags(self):
        """Détecter red flags investisseurs"""
        console.print("\n[cyan]🚩 DÉTECTION RED FLAGS[/]")

        red_flags = self.rules['red_flags']

        # Check CAC/LTV ratio
        # Note: Ces données doivent être dans assumptions pour être vérifiées
        if 'cac' in self.assumptions.get('sales_assumptions', {}).get('enterprise_hub', {}):
            cac = self.assumptions['sales_assumptions']['enterprise_hub']['cac']
            ltv = self.assumptions['sales_assumptions']['enterprise_hub'].get('ltv', 0)

            if ltv > 0 and cac > ltv / 3:
                self.warnings.append({
                    'type': 'RED FLAG',
                    'message': f"CAC {cac:,}€ > LTV/3 ({ltv/3:,.0f}€)",
                    'severity': 'CRITICAL'
                })
                console.print(f"  [red]🚩 CAC/LTV ratio défavorable: {cac:,}€ > {ltv/3:,.0f}€[/]")
            else:
                console.print(f"  [green]✓ CAC/LTV ratio sain[/]")

        # Check churn
        churn = self.assumptions['pricing']['enterprise_hub']['churn_annual']
        max_churn = red_flags['churn_annual']['max']

        if churn > max_churn:
            self.errors.append({
                'type': 'RED FLAG',
                'message': f"Churn {churn:.1%} > {max_churn:.0%}",
                'severity': 'CRITICAL'
            })
            console.print(f"  [red]🚩 Churn trop élevé: {churn:.1%}[/]")
        else:
            console.print(f"  [green]✓ Churn acceptable: {churn:.1%}[/]")

    def check_error_patterns(self, doc_path: Path):
        """Vérifier patterns d'erreurs fréquentes"""
        console.print("\n[cyan]🔎 VÉRIFICATION PATTERNS D'ERREURS[/]")

        doc = Document(doc_path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])

        patterns = self.rules['error_patterns']

        for pattern_rule in patterns:
            pattern = pattern_rule['pattern']
            matches = re.findall(pattern, full_text, re.IGNORECASE)

            if matches:
                severity = pattern_rule['severity']
                message = pattern_rule.get('message', f"Pattern {pattern} détecté")

                if severity == 'CRITICAL':
                    self.errors.append({
                        'type': 'PATTERN ERREUR',
                        'pattern': pattern,
                        'occurrences': len(matches),
                        'message': message
                    })
                    console.print(f"  [red]✗ Pattern critique détecté: {pattern} ({len(matches)}x)[/]")
                else:
                    self.warnings.append({
                        'type': 'PATTERN WARNING',
                        'pattern': pattern,
                        'occurrences': len(matches),
                        'message': message
                    })
                    console.print(f"  [yellow]⚠ Pattern warning: {pattern} ({len(matches)}x)[/]")

    def generate_report(self) -> Dict[str, Any]:
        """Générer rapport de cohérence"""
        console.print("\n" + "=" * 60)
        console.print("[bold]📊 RAPPORT DE COHÉRENCE[/]")
        console.print("=" * 60)

        # Créer table des erreurs
        if self.errors:
            error_table = Table(title="❌ ERREURS CRITIQUES", show_header=True, header_style="bold red")
            error_table.add_column("Section", style="cyan")
            error_table.add_column("Problème", style="white")
            error_table.add_column("Impact", justify="center")

            for error in self.errors:
                section = error.get('section', 'N/A')
                probleme = error.get('probleme', error.get('message', 'N/A'))
                impact = f"{error.get('impact', '?')}/5" if 'impact' in error else error.get('severity', 'N/A')
                error_table.add_row(section, probleme, str(impact))

            console.print(error_table)

        # Créer table des corrections
        if self.corrections:
            console.print("\n[bold green]✅ CORRECTIONS PROPOSÉES[/]")

            for correction in self.corrections:
                panel_content = f"""[bold]Section:[/] {correction['section']}
[bold]Champ:[/] {correction['field']}

[bold cyan]Options:[/]
"""
                for opt_name, opt_data in correction['options'].items():
                    panel_content += f"\n[bold]{opt_name.upper()}:[/] {opt_data['value']} ({opt_data['multiple']}x ARR)\n  → {opt_data['justification']}\n"

                panel_content += f"\n[bold green]RECOMMANDATION:[/] {correction['recommendation'].upper()}"

                console.print(Panel(panel_content, border_style="green"))

        # Warnings
        if self.warnings:
            console.print(f"\n[yellow]⚠️  {len(self.warnings)} WARNINGS[/]")
            for warning in self.warnings[:5]:
                console.print(f"  • {warning.get('message', warning.get('type'))}")

        # Statut final
        status = "❌ ÉCHEC" if self.errors else "✅ SUCCÈS"
        console.print(f"\n[bold]Statut:[/] {status}")
        console.print(f"Erreurs: {len(self.errors)}, Warnings: {len(self.warnings)}")

        return {
            'status': 'FAILED' if self.errors else 'PASSED',
            'errors': self.errors,
            'warnings': self.warnings,
            'corrections': self.corrections,
            'timestamp': datetime.now().isoformat()
        }


def main():
    """Fonction principale"""
    logger.info("=" * 60)
    logger.info("🚀 VALIDATION COHÉRENCE - GenieFactory BP 14 Mois")
    logger.info("=" * 60)

    base_path = Path(__file__).parent.parent

    # Charger données
    projections_path = base_path / "data" / "structured" / "projections.json"
    assumptions_path = base_path / "data" / "structured" / "assumptions.yaml"
    rules_path = base_path / "data" / "validation_rules.yaml"
    word_path = base_path / "data" / "outputs" / "BM_Updated_14M.docx"

    with open(projections_path, 'r', encoding='utf-8') as f:
        projections = json.load(f)

    with open(assumptions_path, 'r', encoding='utf-8') as f:
        assumptions = yaml.safe_load(f)

    with open(rules_path, 'r', encoding='utf-8') as f:
        rules = yaml.safe_load(f)

    # Valider
    validator = CoherenceValidator(projections, assumptions, rules)
    report = validator.validate_all(word_path)

    # Sauvegarder rapport
    logs_dir = base_path / "logs"
    logs_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = logs_dir / f"coherence_report_{timestamp}.txt"

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(f"RAPPORT VALIDATION COHÉRENCE\n")
        f.write(f"Date: {report['timestamp']}\n")
        f.write(f"Status: {report['status']}\n\n")
        f.write(f"Erreurs: {len(report['errors'])}\n")
        f.write(f"Warnings: {len(report['warnings'])}\n\n")

        if report['errors']:
            f.write("ERREURS:\n")
            for error in report['errors']:
                f.write(f"- {error}\n")

        if report['corrections']:
            f.write("\nCORRECTIONS PROPOSÉES:\n")
            for correction in report['corrections']:
                f.write(f"- {correction}\n")

    # Sauvegarder corrections en YAML
    if report['corrections']:
        corrections_path = base_path / "data" / "structured" / "corrections_proposed.yaml"
        with open(corrections_path, 'w', encoding='utf-8') as f:
            yaml.dump(report['corrections'], f, allow_unicode=True, default_flow_style=False)
        logger.info(f"📄 Corrections proposées: {corrections_path}")

    logger.info(f"📄 Rapport sauvegardé: {report_path}")

    return 0 if report['status'] == 'PASSED' else 1


if __name__ == "__main__":
    exit(main())
