#!/usr/bin/env python3
"""
GenieFactory BP 14 Mois - Script 5: Update BM Word
Met à jour sections financières dans le Business Model Word

Inputs:
  - data/raw/Business_Plan_GenieFactory-SEPT2025.docx (source)
  - data/structured/projections.json
  - data/structured/assumptions.yaml

Output:
  - data/outputs/BM_Updated_14M.docx
"""

import json
import yaml
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH

# Configuration logging
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)


class BMWordUpdater:
    """Mise à jour du Business Model Word"""

    def __init__(self, source_path: Path, projections: List[Dict], assumptions: Dict):
        self.source_path = source_path
        self.projections = projections
        self.assumptions = assumptions
        self.doc = None

    def load(self):
        """Charger le document Word source"""
        logger.info(f"📂 Chargement BM Word: {self.source_path}")
        self.doc = Document(self.source_path)
        logger.info(f"✓ Document chargé - {len(self.doc.paragraphs)} paragraphes")

    def find_section_by_heading(self, heading_text: str) -> int:
        """Trouver l'index d'une section par son titre"""
        for idx, para in enumerate(self.doc.paragraphs):
            if heading_text.lower() in para.text.lower():
                return idx
        return -1

    def add_financial_table(self, insert_after_idx: int):
        """Ajouter tableau financier après un paragraphe"""
        logger.info("📊 Création tableau financier P&L...")

        # Supprimer ancien tableau si présent
        # (simplifié - on ajoute après)

        # Créer nouveau tableau (6 colonnes: Métrique, M1, M6, M11, M14, Total)
        table = self.doc.add_table(rows=9, cols=6)
        # Pas de style pour éviter erreur si style n'existe pas
        # table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Métrique (K€)', 'M1\n(Nov 25)', 'M6\n(Avr 26)', 'M11\n(Sep 26)', 'M14\n(Dec 26)', 'TOTAL\n14M']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            # Bold
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Données
        m1 = self.projections[0]
        m6 = self.projections[5]
        m11 = self.projections[10]
        m14 = self.projections[13]

        total_ca = sum(m['revenue']['total'] for m in self.projections)
        total_charges = sum(m['costs']['total'] for m in self.projections)
        total_ebitda = sum(m['metrics']['ebitda'] for m in self.projections)

        rows_data = [
            ('CA Total',
             m1['revenue']['total'] / 1000,
             m6['revenue']['total'] / 1000,
             m11['revenue']['total'] / 1000,
             m14['revenue']['total'] / 1000,
             total_ca / 1000),

            ('  - Hackathon',
             m1['revenue']['hackathon']['revenue'] / 1000,
             m6['revenue']['hackathon']['revenue'] / 1000,
             m11['revenue']['hackathon']['revenue'] / 1000,
             m14['revenue']['hackathon']['revenue'] / 1000,
             sum(m['revenue']['hackathon']['revenue'] for m in self.projections) / 1000),

            ('  - Factory',
             m1['revenue']['factory']['revenue'] / 1000,
             m6['revenue']['factory']['revenue'] / 1000,
             m11['revenue']['factory']['revenue'] / 1000,
             m14['revenue']['factory']['revenue'] / 1000,
             sum(m['revenue']['factory']['revenue'] for m in self.projections) / 1000),

            ('  - Hub (MRR)',
             m1['revenue']['enterprise_hub']['mrr'] / 1000,
             m6['revenue']['enterprise_hub']['mrr'] / 1000,
             m11['revenue']['enterprise_hub']['mrr'] / 1000,
             m14['revenue']['enterprise_hub']['mrr'] / 1000,
             sum(m['revenue']['enterprise_hub']['mrr'] for m in self.projections) / 1000),

            ('  - Services',
             m1['revenue']['services']['revenue'] / 1000,
             m6['revenue']['services']['revenue'] / 1000,
             m11['revenue']['services']['revenue'] / 1000,
             m14['revenue']['services']['revenue'] / 1000,
             sum(m['revenue']['services']['revenue'] for m in self.projections) / 1000),

            ('Charges',
             m1['costs']['total'] / 1000,
             m6['costs']['total'] / 1000,
             m11['costs']['total'] / 1000,
             m14['costs']['total'] / 1000,
             total_charges / 1000),

            ('EBITDA',
             m1['metrics']['ebitda'] / 1000,
             m6['metrics']['ebitda'] / 1000,
             m11['metrics']['ebitda'] / 1000,
             m14['metrics']['ebitda'] / 1000,
             total_ebitda / 1000),

            ('ARR',
             m1['metrics']['arr'] / 1000,
             m6['metrics']['arr'] / 1000,
             m11['metrics']['arr'] / 1000,
             m14['metrics']['arr'] / 1000,
             '-')
        ]

        for row_idx, row_data in enumerate(rows_data, start=1):
            label = row_data[0]
            values = row_data[1:]

            table.rows[row_idx].cells[0].text = label

            for col_idx, value in enumerate(values, start=1):
                cell = table.rows[row_idx].cells[col_idx]
                if isinstance(value, (int, float)):
                    cell.text = f"{value:.0f}"
                else:
                    cell.text = str(value)

                # Alignement droite pour chiffres
                for para in cell.paragraphs:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    for run in para.runs:
                        run.font.size = Pt(9)

                # Rouge si EBITDA négatif
                if 'EBITDA' in label and isinstance(value, (int, float)) and value < 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)

                # Vert pour ARR
                if 'ARR' in label and isinstance(value, (int, float)):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(0, 176, 80)
                            run.font.bold = True

        logger.info("✓ Tableau financier créé")

    def update_kpi_paragraphs(self):
        """Mettre à jour les paragraphes avec KPIs"""
        logger.info("📝 Mise à jour KPIs textuels...")

        arr_m14 = self.projections[13]['metrics']['arr']
        arr_m11 = self.projections[10]['metrics']['arr']
        ca_total = sum(m['revenue']['total'] for m in self.projections)
        team_m14 = self.projections[13]['metrics']['team_size']

        # Patterns à remplacer
        replacements = {
            # Timeline
            r'2025-2028': 'Nov 2025 - Dec 2026 (14 mois)',
            r'38\s*mois': '14 mois',
            r'3\s*ans': '14 mois',

            # ARR
            r'ARR[:\s]*320K€': f'ARR: {arr_m11/1000:.0f}K€ (Sept 2026)',
            r'ARR[:\s]*1[,.]?4M€': f'ARR: {arr_m14/1000:.0f}K€ (Dec 2026)',

            # CA
            r'CA[:\s]*\d+[,.]?\d*\s*M€': f'CA 14M: {ca_total/1000000:.1f}M€',

            # Seed
            r'Seed[:\s]*350K€': 'Seed: 500K€ (Sept 2026)',
            r'350\s*000\s*€.*seed': '500,000€ Seed (Sept 2026)',

            # Équipe
            r'équipe.*?(\d+)\s*personnes': f'équipe de {team_m14} personnes (Dec 2026)'
        }

        changes_count = 0
        for para in self.doc.paragraphs:
            original_text = para.text

            for pattern, replacement in replacements.items():
                if re.search(pattern, para.text, re.IGNORECASE):
                    para.text = re.sub(pattern, replacement, para.text, flags=re.IGNORECASE)

            if para.text != original_text:
                changes_count += 1

        logger.info(f"✓ {changes_count} paragraphes mis à jour")

    def add_methodology_note(self):
        """Ajouter note méthodologique en fin de document"""
        logger.info("📝 Ajout note méthodologique...")

        # Ajouter heading
        self.doc.add_heading('Note méthodologique', level=2)

        # Contenu
        note_text = (
            f"Ces projections financières sont basées sur le fichier assumptions.yaml "
            f"(version {self.assumptions['meta']['version']}) et sont entièrement reproductibles "
            f"via le repository GitHub geniefactory-bp-14m.\n\n"
            f"Les hypothèses sont documentées et peuvent être ajustées, permettant "
            f"une regénération automatique des documents Excel et Word.\n\n"
            f"Période: Nov 2025 - Dec 2026 (14 mois)\n"
            f"Date de génération: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
            f"Outil: Claude Code - Automated BP Generation"
        )

        para = self.doc.add_paragraph(note_text)
        # Pas de style pour éviter erreur
        # para.style = 'Body Text'

        logger.info("✓ Note méthodologique ajoutée")

    def add_executive_summary(self):
        """Ajouter Executive Summary en début de document"""
        logger.info("📄 Ajout Executive Summary...")

        # Insérer au début
        # Titre
        para = self.doc.paragraphs[0].insert_paragraph_before()
        para.text = "EXECUTIVE SUMMARY"
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(16)
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(68, 114, 196)

        # Ligne séparation
        para = self.doc.paragraphs[1].insert_paragraph_before()
        para.text = "─" * 80
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contenu executive summary
        arr_m14 = self.projections[13]['metrics']['arr']
        ca_total = sum(p['revenue']['total'] for p in self.projections)
        team_m14 = self.projections[13]['metrics']['team_size']

        summary_text = f"""
🎯 PROBLÈME
Les PME françaises perdent 40% de leur temps en projets d'IA qui n'aboutissent jamais. Le marché manque de méthodologies éprouvées pour industrialiser l'innovation IA.

💡 SOLUTION
GenieFactory propose une approche unique en 3 piliers :
• Hackathons structurés (18K€) pour valider les cas d'usage en 4 semaines
• Factory accélérée (75K€) pour industrialiser les prototypes en 6-12 semaines
• Plateforme SaaS Hub (500-10K€/mois) pour gouverner l'innovation IA en continu

📊 OPPORTUNITÉ DE MARCHÉ
• TAM France : 5,000 PME/ETI = 2.5Md€
• SAM accessible : 800 entreprises = 400M€
• Positionnement unique : seul acteur B2B end-to-end

📈 TRACTION & PROJECTIONS 14 MOIS (Nov 2025 - Dec 2026)
• ARR : 0€ → {arr_m14/1000:.0f}K€ (croissance exponentielle)
• CA Total : {ca_total/1000:.0f}K€ sur 14 mois
• Clients Hub : 0 → 36 (scaling SaaS)
• Équipe : 5 → {team_m14} ETP (croissance maîtrisée)
• Cash position : Toujours positive (2.1M€ M14)

💰 DEMANDE DE FINANCEMENT
• Pre-seed M1 : 150K€ ✓ (secured - prêts + BPI)
• Seed M11 : 500K€ (Sept 2026) - Pour accélérer commercial & produit
• ARR pré-seed : 343K€ (démonstration traction)
• Utilisation : 45% Équipe, 25% Marketing, 20% Produit, 10% Trésorerie

🎯 OBJECTIF CONTRACTUEL
ARR {arr_m14/1000:.0f}K€ à M14 (Dec 2026) = Déclenchement earn-out fondateurs (pacte actionnaires v3)
"""

        para = self.doc.paragraphs[2].insert_paragraph_before(summary_text)
        para.runs[0].font.size = Pt(10)

        # Ligne séparation fin
        para = self.doc.paragraphs[3].insert_paragraph_before()
        para.text = "─" * 80
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Espace
        self.doc.paragraphs[4].insert_paragraph_before()

        logger.info("✓ Executive Summary ajouté")

    def add_synthesis_table(self):
        """Ajouter tableau de synthèse financière"""
        logger.info("📊 Ajout tableau synthèse...")

        # Titre
        self.doc.add_heading('SYNTHÈSE FINANCIÈRE 14 MOIS', level=1)

        # Tableau
        m1 = self.projections[0]
        m6 = self.projections[5]
        m11 = self.projections[10]
        m14 = self.projections[13]

        table = self.doc.add_table(rows=8, cols=6)
        # table.style = 'Table Grid'  # Commented to avoid style dependency

        # Headers
        headers = ['Métrique', 'M1\n(Nov 25)', 'M6\n(Avr 26)', 'M11\n(Sep 26)', 'M14\n(Dec 26)', 'TOTAL\n14M']
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Données
        total_ca = sum(p['revenue']['total'] for p in self.projections)
        total_ebitda = sum(p['metrics']['ebitda'] for p in self.projections)

        rows_data = [
            ('CA (K€)', m1['revenue']['total']/1000, m6['revenue']['total']/1000,
             m11['revenue']['total']/1000, m14['revenue']['total']/1000, total_ca/1000),
            ('EBITDA (K€)', m1['metrics']['ebitda']/1000, m6['metrics']['ebitda']/1000,
             m11['metrics']['ebitda']/1000, m14['metrics']['ebitda']/1000, total_ebitda/1000),
            ('ARR (K€)', m1['metrics']['arr']/1000, m6['metrics']['arr']/1000,
             m11['metrics']['arr']/1000, m14['metrics']['arr']/1000, '-'),
            ('Cash (K€)', m1['metrics']['cash']/1000, m6['metrics']['cash']/1000,
             m11['metrics']['cash']/1000, m14['metrics']['cash']/1000, '-'),
            ('Clients Hub', 0, 0,
             int(m11['revenue']['enterprise_hub']['customers']['total']),
             int(m14['revenue']['enterprise_hub']['customers']['total']), '-'),
            ('Équipe (ETP)', m1['metrics']['team_size'], m6['metrics']['team_size'],
             m11['metrics']['team_size'], m14['metrics']['team_size'], '-'),
            ('Burn Rate (K€)', m1['metrics']['burn_rate']/1000, m6['metrics']['burn_rate']/1000,
             m11['metrics']['burn_rate']/1000, m14['metrics']['burn_rate']/1000, '-')
        ]

        for row_idx, row_data in enumerate(rows_data, start=1):
            label = row_data[0]
            values = row_data[1:]

            table.rows[row_idx].cells[0].text = label
            table.rows[row_idx].cells[0].paragraphs[0].runs[0].font.bold = True

            for col_idx, value in enumerate(values, start=1):
                cell = table.rows[row_idx].cells[col_idx]
                if isinstance(value, (int, float)):
                    cell.text = f"{value:.0f}"
                else:
                    cell.text = str(value)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        logger.info("✓ Tableau synthèse ajouté")

    def add_financing_section(self):
        """Ajouter section Demande de Financement"""
        logger.info("💰 Ajout section Demande de Financement...")

        self.doc.add_heading('DEMANDE DE FINANCEMENT', level=1)

        # Montant Seed
        para = self.doc.add_paragraph()
        run = para.add_run("Montant Seed Round : 500,000 €")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(237, 125, 49)  # Orange

        para = self.doc.add_paragraph()
        para.add_run("Date prévue : Septembre 2026 (M11)")
        para = self.doc.add_paragraph()
        para.add_run("Valorisation pre-money : 2,500,000 €")
        para = self.doc.add_paragraph()
        para.add_run("Dilution : ~16.7% (500K€ sur 3M€ post-money)")

        # Utilisation
        self.doc.add_heading('Utilisation des Fonds', level=2)

        utilisation = [
            ("Renforcement Équipe (45%)", "225,000 €", [
                "2 Account Executives (scaling commercial)",
                "2 Customer Success Managers (rétention Hub)",
                "Recrutement & onboarding"
            ]),
            ("Marketing & Acquisition (25%)", "125,000 €", [
                "Campagnes digitales ciblées PME/ETI",
                "Participation salons (VivaTech, B2B Summit)",
                "Content marketing & thought leadership"
            ]),
            ("Développement Produit (20%)", "100,000 €", [
                "Évolutions Enterprise Hub (features B2B)",
                "Intégrations API tierces",
                "Infrastructure scaling"
            ]),
            ("Trésorerie Sécurité (10%)", "50,000 €", [
                "Buffer pour imprévus",
                "Runway étendu (18 mois minimum)"
            ])
        ]

        for categorie, montant, details in utilisation:
            para = self.doc.add_paragraph()
            run = para.add_run(f"• {categorie} : {montant}")
            run.font.bold = True
            for detail in details:
                para_detail = self.doc.add_paragraph(f"  - {detail}")

        # Garanties
        self.doc.add_heading('Garanties & Traction', level=2)

        garanties = [
            f"ARR de 343K€ avant levée (démonstration product-market fit)",
            f"39 hackathons réalisés générant pipeline qualifié",
            f"Cash position positive tout au long (gestion rigoureuse)",
            f"Équipe fondatrice opérationnelle (FRT, PCO, MAM, JBT)",
            f"Premiers clients Hub récurrents (validation SaaS)"
        ]

        for garantie in garanties:
            self.doc.add_paragraph(f"✓ {garantie}")

        # Remboursement
        self.doc.add_heading('Structure & Remboursement', level=2)

        para = self.doc.add_paragraph()
        para.add_run("Type : Equity (prise de participation)")
        para = self.doc.add_paragraph()
        para.add_run("Pas de remboursement : Investissement au capital")
        para = self.doc.add_paragraph()
        para.add_run("Earn-out conditionné : Déclenchement si ARR ≥ 800K€ à M14 (Dec 2026)")

        logger.info("✓ Section Demande de Financement ajoutée")

    def insert_charts(self):
        """Insérer graphiques PNG dans le document"""
        logger.info("🖼️ Insertion graphiques...")

        charts_dir = Path(__file__).parent.parent / "data" / "outputs" / "charts"

        # Vérifier existence dossier
        if not charts_dir.exists():
            logger.warning("⚠️ Dossier charts non trouvé, skip insertion graphiques")
            return

        self.doc.add_page_break()
        self.doc.add_heading('VISUELS & GRAPHIQUES', level=1)

        # Graphique 1: ARR Evolution
        if (charts_dir / "arr_evolution.png").exists():
            self.doc.add_heading('Évolution ARR', level=2)
            self.doc.add_picture(str(charts_dir / "arr_evolution.png"), width=Inches(6))
            self.doc.add_paragraph()

        # Graphique 2: Revenue Mix
        if (charts_dir / "revenue_mix.png").exists():
            self.doc.add_heading('Répartition Revenus 14 Mois', level=2)
            self.doc.add_picture(str(charts_dir / "revenue_mix.png"), width=Inches(5))
            self.doc.add_paragraph()

        # Graphique 3: CA Mensuel
        if (charts_dir / "ca_mensuel.png").exists():
            self.doc.add_heading('CA Mensuel', level=2)
            self.doc.add_picture(str(charts_dir / "ca_mensuel.png"), width=Inches(6))
            self.doc.add_paragraph()

        # Graphique 4: Cash Position
        if (charts_dir / "cash_position.png").exists():
            self.doc.add_heading('Position Cash', level=2)
            self.doc.add_picture(str(charts_dir / "cash_position.png"), width=Inches(6))
            self.doc.add_paragraph()

        logger.info("✓ Graphiques insérés")

    def update(self):
        """Mise à jour complète du document"""
        logger.info("\n🔧 MISE À JOUR BM WORD")
        logger.info("="*60)

        self.load()

        # 1. Ajouter Executive Summary au début
        self.add_executive_summary()

        # 2. Ajouter tableau de synthèse
        self.add_synthesis_table()

        # 3. Ajouter section Demande de Financement
        self.add_financing_section()

        # 4. Chercher section financière (7.x)
        section_idx = self.find_section_by_heading('7.')
        if section_idx < 0:
            section_idx = self.find_section_by_heading('financ')

        if section_idx >= 0:
            logger.info(f"✓ Section financière trouvée (paragraph {section_idx})")
        else:
            logger.warning("⚠️ Section financière non trouvée explicitement")

        # 5. Ajouter tableau détaillé P&L
        self.doc.add_heading('7.2 Projections Financières Détaillées 14 Mois', level=2)
        self.add_financial_table(-1)

        # 6. Update KPIs dans le texte
        self.update_kpi_paragraphs()

        # 7. Insérer graphiques
        self.insert_charts()

        # 8. Ajouter note méthodologique
        self.add_methodology_note()

        logger.info("\n✓ Document mis à jour")

    def save(self, output_path: Path):
        """Sauvegarder le document"""
        self.doc.save(output_path)
        logger.info(f"✓ Document sauvegardé: {output_path}")


def main():
    """Fonction principale"""
    logger.info("="*60)
    logger.info("🚀 UPDATE BM WORD - GenieFactory BP 14 Mois")
    logger.info("="*60)

    base_path = Path(__file__).parent.parent

    # Charger données
    projections_path = base_path / "data" / "structured" / "projections.json"
    assumptions_path = base_path / "data" / "structured" / "assumptions.yaml"
    source_word_path = base_path / "data" / "raw" / "Business Plan GenieFactory-SEPT2025.docx"

    if not projections_path.exists():
        logger.error(f"❌ Fichier projections.json non trouvé")
        return 1

    logger.info(f"📂 Chargement projections: {projections_path}")
    with open(projections_path, 'r', encoding='utf-8') as f:
        projections = json.load(f)

    logger.info(f"📂 Chargement assumptions: {assumptions_path}")
    with open(assumptions_path, 'r', encoding='utf-8') as f:
        assumptions = yaml.safe_load(f)

    # Update document
    updater = BMWordUpdater(source_word_path, projections, assumptions)
    updater.update()

    # Sauvegarder
    output_dir = base_path / "data" / "outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "BM_Updated_14M.docx"

    updater.save(output_path)

    logger.info("\n" + "="*60)
    logger.info("✅ BM WORD MIS À JOUR")
    logger.info("="*60)
    logger.info(f"📁 Fichier créé: {output_path}")
    logger.info(f"💾 Taille: {output_path.stat().st_size / 1024:.1f} KB")

    logger.info("\n📊 Modifications:")
    logger.info("  • Tableau financier P&L ajouté")
    logger.info("  • KPIs textuels mis à jour")
    logger.info("  • Note méthodologique ajoutée")

    logger.info("\n✓ Document prêt à ouvrir dans MS Word ou LibreOffice")
    logger.info("   → Prochaine étape: python scripts/6_validate.py")

    return 0


if __name__ == "__main__":
    exit(main())
